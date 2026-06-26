import os
from jinja2 import Template
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx import Document
from models.report import ResearchReport

DEFAULT_MARKDOWN_TEMPLATE = """# Research Gap Analysis Report

## Executive Summary
{{ report.executive_summary }}

## Method Comparison
{{ report.method_comparison }}

## Trend Analysis
{{ report.trend_analysis }}

## Research Gaps
{{ report.research_gaps }}

## Innovation Opportunities
{{ report.innovation_opportunities }}

## References
{% for ref in report.references %}
- {{ ref }}
{% endfor %}
"""

class ReportExporter:
    """
    Exports the structured ResearchReport into Markdown, PDF, and DOCX formats.
    Supports Jinja2 templating for the Markdown base.
    """
    def __init__(self, template_str: str = DEFAULT_MARKDOWN_TEMPLATE):
        self.template_str = template_str

    def export_markdown(self, report: ResearchReport, output_path: str):
        template = Template(self.template_str)
        rendered_md = template.render(report=report)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(rendered_md)
        return output_path

    def export_pdf(self, report: ResearchReport, output_path: str):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Helper to add section
        def add_section(title, content):
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(pdf.epw, 10, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 12)
            pdf.multi_cell(pdf.epw, 7, content)
            pdf.ln(5)

        pdf.set_font("Helvetica", "B", 24)
        pdf.cell(pdf.epw, 15, "Research Gap Analysis Report", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(10)

        add_section("Executive Summary", report.executive_summary)
        add_section("Method Comparison", report.method_comparison)
        add_section("Trend Analysis", report.trend_analysis)
        add_section("Research Gaps", report.research_gaps)
        add_section("Innovation Opportunities", report.innovation_opportunities)

        # References
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(pdf.epw, 10, "References", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 12)
        for idx, ref in enumerate(report.references):
            pdf.multi_cell(pdf.epw, 7, f"[{idx+1}] {ref}")
            
        pdf.output(output_path)
        return output_path

    def export_docx(self, report: ResearchReport, output_path: str):
        doc = Document()
        doc.add_heading("Research Gap Analysis Report", 0)

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary)

        doc.add_heading("Method Comparison", level=1)
        doc.add_paragraph(report.method_comparison)

        doc.add_heading("Trend Analysis", level=1)
        doc.add_paragraph(report.trend_analysis)

        doc.add_heading("Research Gaps", level=1)
        doc.add_paragraph(report.research_gaps)

        doc.add_heading("Innovation Opportunities", level=1)
        doc.add_paragraph(report.innovation_opportunities)

        doc.add_heading("References", level=1)
        for idx, ref in enumerate(report.references):
            doc.add_paragraph(f"[{idx+1}] {ref}", style='List Bullet')

        doc.save(output_path)
        return output_path
